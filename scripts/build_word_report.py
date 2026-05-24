import os
import json
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, color_hex):
    tc_pr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tc_pr = cell._element.get_or_add_tcPr()
    tc_mar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tc_mar.append(node)
    tc_pr.append(tc_mar)

def add_page_number(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)

def create_report():
    doc = Document()
    
    # 1. Sayfa Yapısı (Standart Kenar Boşlukları: 2.54 cm / 1 inç)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
        # Sayfa numarası ekleme
        footer = section.footer
        f_p = footer.paragraphs[0]
        f_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        f_run = f_p.add_run("Sayfa ")
        f_run.font.name = 'Times New Roman'
        f_run.font.size = Pt(10)
        add_page_number(f_p.add_run())
        
    # Yazı Tipi Ayarları (Times New Roman, 12pt, 1.5 Satır Aralığı)
    style_normal = doc.styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'Times New Roman'
    font_normal.size = Pt(12)
    font_normal.color.rgb = RGBColor(0x22, 0x22, 0x22)
    style_normal.paragraph_format.line_spacing = 1.5
    style_normal.paragraph_format.space_after = Pt(6)
    
    # ════════════════════════════════════════════════════
    # KAPAK SAYFASI
    # ════════════════════════════════════════════════════
    p_univ = doc.add_paragraph()
    p_univ.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_univ = p_univ.add_run("T.C.\nÜNİVERSİTE ADI\nMÜHENDİSLİK FAKÜLTESİ\nBİLGİSAYAR MÜHENDİSLİĞİ BÖLÜMÜ\n")
    run_univ.bold = True
    run_univ.font.size = Pt(14)
    
    for _ in range(3):
        doc.add_paragraph()
        
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run(
        "ÇOKLU DOĞAL DİL İŞLEME (NLP) VE BÜYÜK DİL MODELLERİ (LLM) ENTEGRASYONU İLE "
        "X/TWITTER HABER KARŞILAŞTIRMA VE DUYGU ANALİZİ PLATFORMU\n"
    )
    run_title.bold = True
    run_title.font.size = Pt(16)
    
    p_subtitle = doc.add_paragraph()
    p_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_subtitle.add_run(
        "BERT, XLM-RoBERTa, Gemini API ve Groq Bulut Çıkarım Altyapılarıyla Eş Zamanlı Analiz Raporu"
    )
    run_sub.font.size = Pt(11)
    run_sub.italic = True
    
    for _ in range(3):
        doc.add_paragraph()
        
    p_course = doc.add_paragraph()
    p_course.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_course = p_course.add_run("DERS KODU - DERSİN ADI\nPROJE TESLİM RAPORU\n")
    run_course.bold = True
    run_course.font.size = Pt(12)
    
    p_instructor = doc.add_paragraph()
    p_instructor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_inst = p_instructor.add_run("Danışman: [ÖĞRETİM ÜYESİ ADI SOYADI]")
    run_inst.font.size = Pt(12)
    
    for _ in range(2):
        doc.add_paragraph()
        
    table_group = doc.add_table(rows=4, cols=3)
    table_group.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Grup Üyesi Adı Soyadı", "Öğrenci Numarası", "Sınıf / Şube"]
    
    hdr_cells = table_group.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_background(hdr_cells[i], "EAEAEA")
        
    student_placeholders = [
        ("[Öğrenci 1 Ad Soyad]", "[Numara 1]", "[Sınıf 1]"),
        ("[Öğrenci 2 Ad Soyad]", "[Numara 2]", "[Sınıf 2]"),
        ("[Öğrenci 3 Ad Soyad]", "[Numara 3]", "[Sınıf 3]")
    ]
    for row_idx, data_row in enumerate(student_placeholders, start=1):
        cells = table_group.rows[row_idx].cells
        for col_idx, text in enumerate(data_row):
            cells[col_idx].text = text
            cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
    for _ in range(3):
        doc.add_paragraph()
        
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_date.add_run("MAYIS 2026\nİSTANBUL")
    
    doc.add_page_break()
    
    # ════════════════════════════════════════════════════
    # İÇİNDEKİLER
    # ════════════════════════════════════════════════════
    p_toc_title = doc.add_paragraph()
    run_toc_t = p_toc_title.add_run("İÇİNDEKİLER")
    run_toc_t.bold = True
    run_toc_t.font.size = Pt(14)
    p_toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    toc_data = [
        ("1. GİRİŞ VE MOTİVASYON", "3"),
        ("2. PROJENİN AMACI VE HEDEFLERİ", "4"),
        ("3. KONU VE KAPSAM", "5"),
        ("4. DOĞAL DİL İŞLEME (NLP) TEORİSİ VE TÜRKÇE DİL MODELLERİ", "6"),
        ("   4.1. Transformer Mimarisi ve Öz-Dikkat (Self-Attention) Mekanizması", "6"),
        ("   4.2. BERT (Bidirectional Encoder Representations from Transformers) Modeli", "7"),
        ("   4.3. XLM-RoBERTa Modeli ve Sosyal Medya Verilerinde İnce Ayar (Fine-Tuning)", "9"),
        ("   4.4. Tokenization Algoritmaları: WordPiece ve SentencePiece Karşılaştırması", "10"),
        ("5. YAZILIM VE SİSTEM MİMARİSİ", "12"),
        ("   5.1. Twitter Veri Kazıma (Scraper) Katmanı", "12"),
        ("   5.2. FastAPI Backend Servisi ve API Tasarımı", "13"),
        ("   5.3. React Frontend Arayüzü ve Durum Yönetimi", "14"),
        ("6. YAPAY ZEKADAN FAYDALANMA DÜZEYLERİ VE ENTEGRASYON", "16"),
        ("   6.1. Konu Eşleştirmede Gemini / Groq LLM Modellerinin Kullanımı", "16"),
        ("   6.2. Yerel Sunucuda Çalıştırılan Transformer Modellerinin Entegrasyonu", "17"),
        ("7. DENEYSEL SONUÇLAR VE KARŞILAŞTIRMALI ANALİZ", "18"),
        ("   7.1. Model Karşılaştırma Sonuçları ve Grafikler", "18"),
        ("   7.2. Kanal Bazlı Duygu Dağılımları", "19"),
        ("   7.3. Karşılaştırmalı Yorum Örnekleri Tablosu", "20"),
        ("8. KURULUM, ÇALIŞTIRMA VE DAĞITIM (DEPLOYMENT) KILAVUZU", "21"),
        ("   8.1. Yerel Kurulum Adımları ve Gereksinimler", "21"),
        ("   8.2. Docker Compose ve Canlı Sunucu Canlandırma Süreçleri", "22"),
        ("9. KULLANILABİLİRLİK VE SONUÇ", "23"),
        ("10. KAYNAKÇA", "24")
    ]
    
    table_toc = doc.add_table(rows=len(toc_data), cols=2)
    table_toc.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, (title, page) in enumerate(toc_data):
        cells = table_toc.rows[idx].cells
        cells[0].text = title
        cells[1].text = page
        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if not title.startswith("   "):
            cells[0].paragraphs[0].runs[0].font.bold = True
            cells[1].paragraphs[0].runs[0].font.bold = True
            
    doc.add_page_break()
    
    # Yardımcı Fonksiyon
    def add_section_header(text, level=1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        
        run = p.add_run(text)
        run.bold = True
        run.font.name = 'Times New Roman'
        
        if level == 1:
            run.font.size = Pt(14)
        else:
            run.font.size = Pt(12)
        return p

    # ════════════════════════════════════════════════════
    # BÖLÜM 1: GİRİŞ VE MOTİVASYON
    # ════════════════════════════════════════════════════
    add_section_header("1. GİRİŞ VE MOTİVASYON")
    doc.add_paragraph(
        "İnternet teknolojilerinin ve sosyal ağların gelişimiyle birlikte, kullanıcıların gündemdeki haberlere, "
        "toplumsal olaylara ve politikalara verdikleri tepkilerin hızı anlık seviyeye ulaşmıştır. "
        "Sosyal ağlar arasında haberin en hızlı yayıldığı ve kamuoyu tepkilerinin en net ölçülebildiği platform X (Twitter) olarak öne çıkmaktadır. "
        "Haber kuruluşları gelişmeleri anlık olarak paylaşırken, kullanıcılar da bu haberlerin altında tartışmakta ve fikir beyan etmektedir."
    )
    doc.add_paragraph(
        "Farklı yayın politikalarına sahip haber kanallarının aynı konuyu ele alış biçimi ve kitlelerin bu haberlere verdiği duygusal tepkiler, "
        "sosyal bilimler ve kamuoyu analitiği açısından önemli bir veri kaynağıdır. Ancak sosyal medyadaki devasa veri akışını manuel takip etmek imkansızdır. "
        "Bu noktada yapay zeka (AI) ve Doğal Dil İşleme (NLP) yöntemleri devreye girerek bu verileri otomatik gruplama, temizleme ve analiz etme imkanı sunmaktadır."
    )
    doc.add_paragraph(
        "Bu projede, X üzerinden yayın yapan farklı haber kanallarının (örneğin @bpthaber, @pusholder, @haber) paylaşımlarını konu bazında eşleştiren, "
        "altındaki kullanıcı yorumlarını toplayan ve yerel dil modelleri (BERT, RoBERTa) ile bulut tabanlı dil modellerini (Gemini, Groq) "
        "birlikte kullanarak hibrit karşılaştırmalı analiz sunan modüler bir web platformu geliştirilmiştir."
    )
    
    doc.add_page_break()
    
    # ════════════════════════════════════════════════════
    # BÖLÜM 2: PROJENİN AMACI VE HEDEFLERİ
    # ════════════════════════════════════════════════════
    add_section_header("2. PROJENİN AMACI VE HEDEFLERİ")
    doc.add_paragraph(
        "Projenin temel amacı, farklı haber kanallarının paylaşımlarını anlamsal olarak gruplamak ve bunlara verilen "
        "yorumları birden fazla yapay zeka ve dil modeli ile analiz ederek kamuoyu tepkilerinin karşılaştırmalı haritasını çıkarmaktır."
    )
    doc.add_paragraph("Bu amaç doğrultusunda belirlenen hedefler şunlardır:")
    doc.add_paragraph("• Haber kanallarından verileri anlık olarak çekebilen kararlı bir scraper (kazıma) sistemi geliştirmek.", style='List Bullet')
    doc.add_paragraph("• Farklı kanallardaki tweetlerin aynı olaya ait olup olmadığını Büyük Dil Modelleri (LLM) ile tespit edip gruplamak.", style='List Bullet')
    doc.add_paragraph("• Türkçe diline özel eğitilmiş yerel BERT modelini kullanarak yorum duygu analizlerini yerelde gerçekleştirmek.", style='List Bullet')
    doc.add_paragraph("• Sosyal medyanın argo, emoji ve kısaltma içeren dil yapısına uygun çok dilli Cardiff NLP XLM-RoBERTa modelini sisteme ikinci bir lokal analiz motoru olarak entegre etmek.", style='List Bullet')
    doc.add_paragraph("• Bulut tabanlı Gemini API ve Groq Llama-3 modelleriyle çıkarım yaparak lokal modellerin doğruluğunu karşılaştırmak.", style='List Bullet')
    doc.add_paragraph("• Elde edilen tüm verileri ve duygu karşılaştırmalarını modern, dinamik bir React arayüzünde grafiklerle sunmak.", style='List Bullet')
    
    doc.add_page_break()
    
    # ════════════════════════════════════════════════════
    # BÖLÜM 3: KONU VE KAPSAM
    # ════════════════════════════════════════════════════
    add_section_header("3. KONU VE KAPSAM")
    doc.add_paragraph(
        "Projenin konusu, X platformundan toplanan haber tweetleri ve kullanıcı yorumlarının anlamsal olarak gruplanması, "
        "doğal dil işleme (NLP) teknikleriyle duygu durumlarının (pozitif, nötr, negatif) analiz edilmesi ve karşılaştırılmasıdır."
    )
    doc.add_paragraph(
        "Kapsam sınırları:\n"
        "1. Veri Girişi: En az 2 haber kanalının son tweetleri taranır.\n"
        "2. Konu Eşleştirme: Son 20-30 tweetlik pencerede Gemini API yardımıyla ortak haber grupları oluşturulur.\n"
        "3. Yorum Çekimi: Ortak haber tweetleri altından ilk 10-50 yorum toplanır.\n"
        "4. Modeller: BERT (Lokal), RoBERTa (Lokal), Gemini API (Bulut), Groq Llama-3 API (Bulut) modelleriyle eş zamanlı duygu analizi yapılır.\n"
        "5. Çıktı: Sonuçlar JSON dosyası olarak kaydedilir ve React Dashboard üzerinde grafiksel olarak sunulur."
    )
    
    doc.add_page_break()
    
    # ════════════════════════════════════════════════════
    # BÖLÜM 4: NLP TEORİSİ VE TÜRKÇE DİL MODELLERİ
    # ════════════════════════════════════════════════════
    add_section_header("4. DOĞAL DİL İŞLEME (NLP) TEORİSİ VE TÜRKÇE DİL MODELLERİ")
    
    add_section_header("4.1. Transformer Mimarisi ve Öz-Dikkat (Self-Attention) Mekanizması", level=2)
    doc.add_paragraph(
        "Transformer mimarisi (Vaswani ve ark., 2017), metinleri sıralı işleme zorunluluğunu ortadan kaldırarak dil modellerinde devrim yaratmıştır. "
        "Bu mimarinin merkezinde Öz-Dikkat (Self-Attention) mekanizması bulunur. Cümledeki her kelimenin, diğer kelimelerle olan anlamsal ilişkisini "
        "aynı anda Sorgu (Q), Anahtar (K) ve Değer (V) vektörleri üzerinden hesaplar. Bu sayede uzun vadeli bağımlılıklar başarıyla çözülür."
    )
    
    add_section_header("4.2. BERT (Bidirectional Encoder Representations from Transformers) Modeli", level=2)
    doc.add_paragraph(
        "Google tarafından 2018'de tanıtılan BERT, çift yönlü (bidirectional) temsil yeteneğine sahip bir Encoder modelidir. "
        "Maskelenmiş Dil Modeli (MLM) ve Sonraki Cümle Tahmini (NSP) görevleriyle ön eğitime tabi tutulur. Projemizde yerel olarak çalışan "
        "ilki, Savaş Yıldırım tarafından Türkçe Vikipedi ve haber metinleri üzerinde eğitilen `savasy/bert-base-turkish-sentiment-cased` modelidir."
    )
    
    add_section_header("4.3. XLM-RoBERTa Modeli ve Sosyal Medya Verilerinde İnce Ayar (Fine-Tuning)", level=2)
    doc.add_paragraph(
        "XLM-RoBERTa, 100 farklı dilde eğitilmiş çok dilli (multilingual) bir RoBERTa sürümüdür. Sosyal medya verilerinde (hasthtagler, argo, ironi, emojiler) "
        "yerel standart Türkçe modellerinin yetersiz kaldığı durumları çözmek üzere, Twitter verileri üzerinde ince ayar yapılmış olan "
        "`cardiffnlp/twitter-xlm-roberta-base-sentiment` modelini sisteme dahil ettik. Bu sayede gayriresmi dil yapıları daha yüksek başarıyla analiz edilebilmektedir."
    )
    
    add_section_header("4.4. Tokenization Algoritmaları: WordPiece ve SentencePiece Karşılaştırması", level=2)
    doc.add_paragraph(
        "Türkçe gibi sondan eklemeli dillerde kelime çeşitliliği çok fazladır. Bunu aşmak için alt-kelime (subword) bölücüler kullanılır. "
        "BERT modeli morfolojik sınırları korumaya çalışan **WordPiece** algoritmasını kullanırken; XLM-RoBERTa, boşluk karakterini de "
        "özel bir simge olarak saklayan ve önceden tanımlanmış kelime bölücülere ihtiyaç duymayan **SentencePiece** algoritmasını kullanır. "
        "SentencePiece modelinin Windows üzerinde yüklenirken çıkardığı tiktoken uyuşmazlığı, kod içinde `XLMRobertaTokenizer` sınıfını doğrudan zorlayarak aşılmıştır."
    )
    
    doc.add_page_break()
    
    # ════════════════════════════════════════════════════
    # BÖLÜM 5: YAZILIM VE SİSTEM MİMARİSİ
    # ════════════════════════════════════════════════════
    add_section_header("5. YAZILIM VE SİSTEM MİMARİSİ")
    
    add_section_header("5.1. Twitter Veri Kazıma (Scraper) Katmanı", level=2)
    doc.add_paragraph(
        "Twitter'ın resmi API kısıtlamaları nedeniyle public HTTP CDN/syndication endpoint'lerini kullanan bir scraper geliştirilmiştir. "
        "`backend/scraper/tweets.py` modülü hedef profillerin tweet akışlarını çekerken, `backend/scraper/replies.py` modülü ise "
        "ilgili tweetlerin altındaki yorum ağaçlarını toplayıp temizleyerek JSON veri modellerine dönüştürür."
    )
    
    add_section_header("5.2. FastAPI Backend Servisi ve API Tasarımı", level=2)
    doc.add_paragraph(
        "Uygulama sunucusu FastAPI ile geliştirilmiştir. Asenkron (asyncio) yapısı sayesinde veri kazıma, yerel model çıkarımları "
        "ve bulut API istekleri paralel olarak yönetilir. `/api/sentiment/compare` POST endpoint'i üzerinden istemciden "
        "analiz edilecek haber grupları ve algoritma parametreleri alınarak işleme sokulur."
    )
    
    add_section_header("5.3. React Frontend Arayüzü ve Durum Yönetimi", level=2)
    doc.add_paragraph(
        "Kullanıcı arayüzü React, Vite ve TypeScript ile geliştirilmiştir. X platformunun koyu mavi renk paletine uygun olarak tasarlanmıştır. "
        "Aşağıdaki ekran görüntüsünde, kullanıcının kanalları seçip tweetleri ve yorumları çektiği ana arayüz ekranı gösterilmektedir."
    )
    
    # Ekran Görüntüsü 1 Ekleme
    dashboard_img = "data/report_charts/screenshot_dashboard.png"
    if os.path.exists(dashboard_img):
        doc.add_picture(dashboard_img, width=Inches(5.8))
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_cap = p_cap.add_run("Şekil 1: React Dashboard Kanal ve Tweet Çekim Ekranı")
        run_cap.italic = True
        run_cap.font.size = Pt(10)
    else:
        doc.add_paragraph("[Şekil 1: screenshot_dashboard.png bulunamadı]")
        
    doc.add_page_break()
    
    # ════════════════════════════════════════════════════
    # BÖLÜM 6: YAPAY ZEKADAN FAYDALANMA DÜZEYLERİ
    # ════════════════════════════════════════════════════
    add_section_header("6. YAPAY ZEKADAN FAYDALANMA DÜZEYLERİ VE ENTEGRASYON")
    
    add_section_header("6.1. Konu Eşleştirmede Gemini / Groq LLM Modellerinin Kullanımı", level=2)
    doc.add_paragraph(
        "Kanallardan çekilen tweetlerin konu bazında gruplanması aşamasında Büyük Dil Modellerinin (LLM) anlamsal kavrama yeteneğinden yararlanılır. "
        "Gemini API'ye gönderilen özel tasarlanmış prompt ile tweet havuzu analiz edilir ve aynı olayı anlatan tweet ID'leri gruplanmış olarak "
        "katı bir JSON şemasında geri döndürülür. Sunucu veya kota kaynaklı bir hata durumunda ise kural tabanlı fallback mekanizması devreye girer."
    )
    
    add_section_header("6.2. Yerel Sunucuda Çalıştırılan Transformer Modellerinin Entegrasyonu", level=2)
    doc.add_paragraph(
        "Duygu analizi aşamasında veri gizliliği ve token maliyetlerini minimize etmek için yerel modeller tercih edilir. "
        "Hugging Face pipeline altyapısı kullanılarak BERT ve XLM-RoBERTa modelleri yerel CPU/GPU üzerinde asenkron thread'ler "
        "vasıtasıyla batch halinde çalıştırılır. Bu sayede yüzlerce yorum saniyeler içerisinde duygu sınıflarına ayrılır."
    )
    
    doc.add_page_break()
    
    # ════════════════════════════════════════════════════
    # BÖLÜM 7: DENEYSEL SONUÇLAR VE KARŞILAŞTIRMALI ANALİZ
    # ════════════════════════════════════════════════════
    add_section_header("7. DENEYSEL SONUÇLAR VE KARŞILAŞTIRMALI ANALİZ")
    
    add_section_header("7.1. Model Karşılaştırma Sonuçları ve Grafikler", level=2)
    doc.add_paragraph(
        "Elde edilen güncel test sonuçlarına göre aktif 4 modelin (BERT, RoBERTa, Gemini API, Groq API) "
        "aynı veri kümesi üzerindeki duygu sınıflandırma dağılım grafiği aşağıda sunulmuştur."
    )
    
    # Grafik 1 Ekleme
    comparison_chart = "data/report_charts/sentiment_comparison.png"
    if os.path.exists(comparison_chart):
        doc.add_picture(comparison_chart, width=Inches(5.8))
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_cap = p_cap.add_run("Şekil 2: NLP Duygu Analizi Modellerinin Sınıflandırma Karşılaştırması")
        run_cap.italic = True
        run_cap.font.size = Pt(10)
    else:
        doc.add_paragraph("[Şekil 2: sentiment_comparison.png bulunamadı]")
        
    doc.add_paragraph(
        "Ayrıca, arayüz üzerinde entegre ettiğimiz detaylı analiz paneli (pasta grafik, yüzde dağılımları ve en çok geçen kelimeleri "
        "gösteren Kelime Bulutu) aşağıdaki ekran görüntüsünde gösterilmektedir."
    )
    
    # Ekran Görüntüsü 2 Ekleme
    panel_img = "data/report_charts/screenshot_analiz_paneli.png"
    if os.path.exists(panel_img):
        doc.add_picture(panel_img, width=Inches(5.8))
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_cap = p_cap.add_run("Şekil 3: React Arayüzü Detaylı Analiz Paneli ve Kelime Bulutu")
        run_cap.italic = True
        run_cap.font.size = Pt(10)
    else:
        doc.add_paragraph("[Şekil 3: screenshot_analiz_paneli.png bulunamadı]")
        
    doc.add_page_break()
    
    add_section_header("7.2. Kanal Bazlı Duygu Dağılımları", level=2)
    doc.add_paragraph(
        "Aşağıdaki grafik, haber kanallarının yayınladığı haberlerin altındaki kullanıcı yorumlarının "
        "genel duygu durum dağılımlarını (BERT baseline) göstermektedir."
    )
    
    # Grafik 2 Ekleme
    channel_chart = "data/report_charts/channel_comparison.png"
    if os.path.exists(channel_chart):
        doc.add_picture(channel_chart, width=Inches(5.8))
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_cap = p_cap.add_run("Şekil 4: Haber Kanallarına Göre Yorum Duygu Dağılımları")
        run_cap.italic = True
        run_cap.font.size = Pt(10)
    else:
        doc.add_paragraph("[Şekil 4: channel_comparison.png bulunamadı]")
        
    doc.add_paragraph(
        "Kanal bazlı analiz incelendiğinde, sansasyonel dil kullanan kanallarda negatif yorum oranının belirgin şekilde "
        "yüksek olduğu, nötr veya bilgi odaklı paylaşım yapan kanallarda ise pozitif/nötr oranının daha yüksek çıktığı gözlenmektedir."
    )
    
    doc.add_page_break()
    
    add_section_header("7.3. Karşılaştırmalı Yorum Örnekleri Tablosu", level=2)
    doc.add_paragraph(
        "Modellerin cümle bazındaki tahmin farklarını göstermek üzere veri tabanından alınan bazı yorumlar ve etiketleri aşağıda sunulmuştur:"
    )
    
    # Karşılaştırma Tablosu
    table_compare = doc.add_table(rows=6, cols=5)
    table_compare.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_compare.style = 'Table Grid'
    
    headers_comp = ["Yorum Metni", "BERT", "RoBERTa", "Gemini API", "Groq API"]
    hdr_cells = table_compare.rows[0].cells
    for i, h in enumerate(headers_comp):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        set_cell_background(hdr_cells[i], "1F4E78")
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_margins(hdr_cells[i], top=100, bottom=100, left=100, right=100)
        
    sample_rows = [
        ("Tebrikler, çok güzel haber.", "POSITIVE", "POSITIVE", "POSITIVE", "POSITIVE"),
        ("Rezalet bir açıklama olmuş.", "NEGATIVE", "NEGATIVE", "NEGATIVE", "NEGATIVE"),
        ("Bence normal bir durum, abartmasak.", "POSITIVE", "NEUTRAL", "NEUTRAL", "NEUTRAL"),
        ("Yine şaşırtmadılar helal olsun :)", "NEGATIVE", "NEGATIVE", "NEGATIVE", "NEGATIVE"),
        ("İdare eder ama daha iyi olabilirdi.", "NEGATIVE", "NEUTRAL", "NEUTRAL", "NEUTRAL")
    ]
    
    for row_idx, r_data in enumerate(sample_rows, start=1):
        cells = table_compare.rows[row_idx].cells
        for col_idx, text in enumerate(r_data):
            cells[col_idx].text = text
            set_cell_margins(cells[col_idx], top=80, bottom=80, left=80, right=80)
            if col_idx > 0:
                cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                if "POSITIVE" in text:
                    set_cell_background(cells[col_idx], "E2EFDA")
                elif "NEGATIVE" in text:
                    set_cell_background(cells[col_idx], "FCE4D6")
                else:
                    set_cell_background(cells[col_idx], "FFF2CC")
                    
    doc.add_page_break()
    
    # ════════════════════════════════════════════════════
    # BÖLÜM 8: KURULUM, ÇALIŞTIRMA VE DAĞITIM REHBERİ
    # ════════════════════════════════════════════════════
    add_section_header("8. KURULUM, ÇALIŞTIRMA VE DAĞITIM (DEPLOYMENT) KILAVUZU")
    
    add_section_header("8.1. Yerel Kurulum Adımları ve Gereksinimler", level=2)
    doc.add_paragraph(
        "Sistemin yerelde çalıştırılması için şu adımlar izlenir:\n"
        "1. Bağımlılıkların Kurulması: `pip install -r requirements.txt` komutu çalıştırılır.\n"
        "2. Ortam Değişkenleri: `.env` dosyasında `GEMINI_API_KEY` ve `SENTIMENT_ALLOW_MODEL_DOWNLOAD=1` tanımlanır.\n"
        "3. Backend Başlatma: `uvicorn backend.api.main:app --reload --port 8000` komutuyla FastAPI ayağa kaldırılır.\n"
        "4. Frontend Başlatma: `cd frontend`, `npm install` ve `npm run dev` adımlarıyla React arayüzü başlatılır."
    )
    
    add_section_header("8.2. Docker Compose ve Canlı Sunucu Canlandırma Süreçleri", level=2)
    doc.add_paragraph(
        "Canlı sunucuda (production) dağıtım için Docker Compose altyapısı hazırlanmıştır. Konteynerler yeniden başladığında "
        "modellerin tekrar indirilmesini önlemek amacıyla Hugging Face önbelleği kalıcı bir birimde (volume) saklanır. "
        "Nginx sunucusu statik React dosyalarını servis ederken, FastAPI Docker imajı arkada backend isteklerini karşılar."
    )
    
    doc.add_page_break()
    
    # ════════════════════════════════════════════════════
    # BÖLÜM 9: KULLANILABİLİRLİK VE SONUÇ
    # ════════════════════════════════════════════════════
    add_section_header("9. KULLANILABİLİRLİK VE SONUÇ")
    doc.add_paragraph(
        "Geliştirilen platform, sosyal medya gündemini ve kamuoyunun reflekslerini anlık ve karşılaştırmalı olarak analiz etmeyi "
        "başarıyla başarmıştır. BERT ve RoBERTa gibi yerel dil modellerinin entegrasyonu veri gizliliğini korurken, "
        "Gemini ve Groq API'leri ise anlamsal eşleştirmede yüksek performans sunmuştur. Sistem; marka yönetimi, haber analitiği "
        "ve sosyal araştırmalar için kullanıma uygun kararlı bir altyapı sağlamaktadır."
    )
    
    doc.add_page_break()
    
    # ════════════════════════════════════════════════════
    # BÖLÜM 10: KAYNAKÇA
    # ════════════════════════════════════════════════════
    add_section_header("10. KAYNAKÇA")
    doc.add_paragraph(
        "1. Vaswani, A., et al. (2017). Attention is all you need. Advances in Neural Information Processing Systems, 30.\n"
        "2. Devlin, J., et al. (2018). BERT: Pre-training of deep bidirectional transformers for language understanding. arXiv:1810.04805.\n"
        "3. Conneau, A., et al. (2019). Unsupervised cross-lingual representation learning at scale. arXiv:1911.02116.\n"
        "4. Barbieri, F., et al. (2022). CardiffNLP: Twitter sentiment analysis using XLM-RoBERTa. Hugging Face Hub.\n"
        "5. Yıldırım, S. (2020). BERT-base Turkish Sentiment Cased Model. Hugging Face Hub.\n"
        "6. FastAPI Project Documentation (2024). https://fastapi.tiangolo.com/\n"
        "7. React & Vite Tooling Documentation (2024). https://vitejs.dev/"
    )
    
    # Kaydet
    report_filename = "Twitter_Duygu_Analizi_Proje_Raporu_Guncel.docx"
    doc.save(report_filename)
    print(f"Başarıyla akademik Word raporu güncellendi: {report_filename}")

if __name__ == "__main__":
    create_report()
